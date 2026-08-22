#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创作追踪数据导出 → 排版优化Word文档
参考《版画实训成长日志完整汇编》的排版风格：
  - 封面标题: 26pt Bold 居中
  - 一级标题(章): 16pt Bold
  - 二级标题(节): 15pt Bold
  - 标签/字段名: 15pt Bold
  - 正文: 11pt
  - 表格: 9pt (保持原有表格内容完整性)
  - 页边距: 左3.2cm 右3.5cm 上下2.5cm
  - A4页面
"""

import pdfplumber
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import unicodedata

PDF_PATH = "E:/橙子文件/2026下乡数据/创作追踪数据导出.pdf"
OUTPUT_PATH = "E:/橙子文件/2026下乡数据/创作追踪数据导出_排版优化.docx"

# ============================================================
# 排版常量（参考《版画实训成长日志完整汇编》）
# ============================================================
FONT_CN = 'DengXian'        # 中文字体（等线）
FONT_EN = 'Arial'            # 英文/数字字体

SIZE_COVER_TITLE = 26        # 封面主标题
SIZE_STUDENT_NAME = 16       # 学生姓名/编号标题 (如 B01)
SIZE_SECTION_TITLE = 16      # 章节副标题 (如 "学生成长日志")
SIZE_CHAPTER = 16            # 章标题
SIZE_SUB_HEADING = 15        # 日期行/标签行 (如 "第 1 天"、"状态标签")
SIZE_BODY = 11               # 正文
SIZE_TABLE = 9               # 表格内容

MARGIN_LEFT = Cm(3.2)        # 左边距 (~90pt)
MARGIN_RIGHT = Cm(3.5)       # 右边距 (~98pt)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)

# 行间距（单倍行距 = 1.0，参考文档大约1.15倍行距）
LINE_SPACING = 1.15
# 段落间距（参考文档约18pt行间距，段落间约16-18pt）
PARA_SPACING_BODY = 4        # 正文段间距
PARA_SPACING_HEADING = 12    # 标题后间距


def sanitize(text):
    """清理非法XML/控制字符"""
    if not text:
        return ''
    result = []
    for ch in str(text):
        cp = ord(ch)
        if cp == 0:
            continue
        if cp < 0x20 and cp not in (0x9, 0xA, 0xD):
            continue
        if 0xD800 <= cp <= 0xDFFF or cp in (0xFFFE, 0xFFFF):
            continue
        if unicodedata.category(ch)[0] == 'C' and ch not in '\t\n\r':
            continue
        result.append(ch)
    return ''.join(result)


def add_run_to_para(para, text, font_cn=FONT_CN, font_en=FONT_EN,
                     size=11, bold=False):
    """向段落添加格式化的文字运行"""
    text = sanitize(text)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    run.bold = bold
    return run


def add_para(doc, text, size=SIZE_BODY, bold=False, align=None,
             after=PARA_SPACING_BODY, font_cn=FONT_CN):
    """添加格式化段落"""
    text = sanitize(text)
    if not text.strip():
        return None
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = Cm(0)
    add_run_to_para(p, text, font_cn=font_cn, size=size, bold=bold)
    return p


def add_label_para(doc, label, value='', size=SIZE_BODY, label_bold=True):
    """添加标签: 值的格式段落（参考文档风格：标签加粗，值正常）"""
    text = sanitize(f"{label}{value}")
    if not text.strip():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(PARA_SPACING_BODY)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = Cm(0)

    if value:
        add_run_to_para(p, label, size=size, bold=label_bold)
        add_run_to_para(p, value, size=size, bold=False)
    else:
        add_run_to_para(p, label, size=size, bold=label_bold)
    return p


def add_separator(doc, size=6):
    """添加分隔空行"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run_to_para(p, '', size=size)
    return p


def add_table_borders(table):
    """为表格添加简洁边框（参考风格：细线边框）"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = OxmlElement('w:tblBorders')
    for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{bn}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '666666')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)

    # 设置表格单元格边距
    tblPr2 = tbl.tblPr
    tbl_cell_mar = OxmlElement('w:tblCellMar')
    for margin_name in ('top', 'left', 'bottom', 'right'):
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '30')
        margin.set(qn('w:type'), 'dxa')
        tbl_cell_mar.append(margin)
    tblPr2.append(tbl_cell_mar)


def set_cell_format(cell, text, size=SIZE_TABLE, bold=False):
    """设置单元格格式"""
    text = sanitize(text)
    for para in cell.paragraphs:
        para.clear()
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.0
    add_run_to_para(para, text, size=size, bold=bold)


def is_data_row(line):
    """B-XX开头且至少5个空格分隔字段"""
    return line.startswith('B-') and len(line.split()) >= 5


def is_table_header(line):
    """判断是否为表头行：多个中文字段用空格分隔"""
    if line.startswith('B-'):
        return False
    parts = line.split()
    if len(parts) < 3:
        return False
    ch_count = sum(1 for p in parts if any('一' <= c <= '鿿' for c in p))
    return ch_count >= len(parts) * 0.5


def is_chapter(line):
    return bool(re.match(r'^第[一二三四五六七八九十]+章\s', line))


def is_section(line):
    return bool(re.match(r'^\d+\.\d+\s+\S', line))


def is_day_title(line):
    return bool(re.match(r'^第\d+天[（(]', line))


def is_student_profile_header(line):
    return bool(re.match(r'^[A-Z]-\d+\s+[\|／]\s', line))


def is_page_number(line):
    return bool(re.match(r'^[（(]\d+[）)]\s*$', line)) or \
           bool(re.match(r'^\d+\s*/\s*\d+\s*$', line))


def is_stats_line(line):
    """统计摘要行"""
    return bool(re.search(r'[：:]\s*\d+\s*(人|个|篇|条|次|%)', line)) or \
           ('|' in line and any(kw in line for kw in
               ('观察', '追踪', '事件', '访谈', '日志', '篇', '人')))


# ============================================================
# 第1步：提取PDF文本
# ============================================================

print("Step 1: Extracting PDF text...")

all_lines = []
page_boundaries = []

with pdfplumber.open(PDF_PATH) as pdf:
    total_pages = len(pdf.pages)
    for pi in range(total_pages):
        page = pdf.pages[pi]
        raw = page.extract_text()
        start = len(all_lines)
        if raw:
            cleaned = raw.replace('\x00', '')
            page_lines = [l.strip() for l in cleaned.split('\n')]
            all_lines.extend(page_lines)
        page_boundaries.append((start, len(all_lines)))

print(f"  {len(all_lines)} lines from {total_pages} pages")

# ============================================================
# 第2步：识别数据表格块
# ============================================================

print("Step 2: Identifying data tables...")

table_blocks = {}
i = 0
while i < len(all_lines):
    line = all_lines[i]
    if is_data_row(line):
        header_idx = i
        for j in range(i - 1, max(i - 4, -1), -1):
            prev = all_lines[j]
            if prev and is_table_header(prev) and not is_data_row(prev):
                header_idx = j
            else:
                break
        data_start = i
        while i < len(all_lines) and is_data_row(all_lines[i]):
            i += 1
        data_end = i - 1
        for k in range(header_idx, data_end + 1):
            table_blocks[k] = (header_idx, data_end)
    else:
        i += 1

block_ranges = []
seen_ends = set()
for k in sorted(table_blocks.keys()):
    hs, de = table_blocks[k]
    if de not in seen_ends:
        block_ranges.append((hs, de))
        seen_ends.add(de)
block_ranges.sort(key=lambda x: x[0])

line_in_block = {}
for hs, de in block_ranges:
    for k in range(hs, de + 1):
        line_in_block[k] = hs

print(f"  {len(block_ranges)} table blocks found")

# ============================================================
# 第3步：创建排版优化的Word文档
# ============================================================

print("Step 3: Building formatted Word document...")

doc = Document()

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = MARGIN_TOP
section.bottom_margin = MARGIN_BOTTOM
section.left_margin = MARGIN_LEFT
section.right_margin = MARGIN_RIGHT

# 默认样式
style = doc.styles['Normal']
style.font.name = FONT_EN
style.font.size = Pt(SIZE_BODY)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
style.paragraph_format.space_after = Pt(PARA_SPACING_BODY)
style.paragraph_format.line_spacing = LINE_SPACING

# ---- 逐页处理 ----
for page_idx in range(total_pages):
    pstart, pend = page_boundaries[page_idx]
    page_lines = all_lines[pstart:pend]

    # ========== 封面页 ==========
    if page_idx == 0:
        for line in page_lines:
            if not line:
                continue
            if any(kw in line for kw in ('创作追踪系统', '数据导出')):
                add_para(doc, line, size=SIZE_COVER_TITLE, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
            elif '自动生成' in line:
                add_para(doc, line, size=SIZE_BODY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
            else:
                # 封面信息行：保持标签+值格式
                if '：' in line or '：' in line:
                    parts = line.split('：', 1) if '：' in line else line.split('：', 1)
                    if len(parts) == 2:
                        add_label_para(doc, parts[0] + '：', parts[1], size=SIZE_BODY)
                    else:
                        add_para(doc, line, size=SIZE_BODY)
                else:
                    add_para(doc, line, size=SIZE_BODY)
        doc.add_page_break()
        continue

    # ========== 目录页 ==========
    if page_idx == 1:
        add_para(doc, '目  录', size=SIZE_CHAPTER, bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
        for line in page_lines:
            if not line or line == '目  录':
                continue
            if is_chapter(line):
                add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=8)
            elif is_section(line):
                add_para(doc, '    ' + line, size=SIZE_BODY, after=3)
            else:
                add_para(doc, line, size=SIZE_BODY, after=3)
        doc.add_page_break()
        continue

    # ========== 内容页 ==========
    processed_blocks = set()
    li = 0

    while li < len(page_lines):
        line = page_lines[li]
        global_li = pstart + li

        # 跳过页码
        if is_page_number(line):
            li += 1
            continue

        # 检查表格块
        if global_li in line_in_block:
            block_start = line_in_block[global_li]
            if block_start not in processed_blocks:
                processed_blocks.add(block_start)
                _, de = table_blocks[block_start]

                # 收集表格数据
                header_row = None
                data_rows = []

                for k in range(block_start, de + 1):
                    parts = all_lines[k].split()
                    if not parts:
                        continue
                    if not all_lines[k].startswith('B-') and is_table_header(all_lines[k]):
                        header_row = parts
                    else:
                        data_rows.append(parts)

                if data_rows:
                    cols = len(header_row) if header_row else max(len(r) for r in data_rows)

                    if header_row:
                        while len(header_row) < cols:
                            header_row.append('')
                        if len(header_row) > cols:
                            header_row = header_row[:cols]

                    uniform_rows = []
                    for r in data_rows:
                        while len(r) < cols:
                            r.append('')
                        if len(r) > cols:
                            r = r[:cols - 1] + [' '.join(r[cols - 1:])]
                        uniform_rows.append(r)

                    total_rows = len(uniform_rows) + (1 if header_row else 0)
                    wt = doc.add_table(rows=total_rows, cols=cols)
                    wt.alignment = WD_TABLE_ALIGNMENT.CENTER
                    add_table_borders(wt)

                    row_offset = 0
                    if header_row:
                        for ci in range(cols):
                            cell = wt.cell(0, ci)
                            set_cell_format(cell,
                                header_row[ci] if ci < len(header_row) else '',
                                size=SIZE_TABLE, bold=True)
                        row_offset = 1

                    for ri, row_data in enumerate(uniform_rows):
                        for ci in range(cols):
                            cell = wt.cell(ri + row_offset, ci)
                            set_cell_format(cell,
                                row_data[ci] if ci < len(row_data) else '',
                                size=SIZE_TABLE)

                    add_separator(doc, size=6)

                li = (de - pstart) + 1
                continue

        # 非表格内容：排版输出
        if not line:
            add_separator(doc, size=4)
            li += 1
            continue

        # ---- 按行类型分级排版 ----
        if is_chapter(line):
            # 第X章 - 一级标题
            add_para(doc, line, size=SIZE_CHAPTER, bold=True, after=14)

        elif is_section(line):
            # X.X - 二级标题
            add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=10)

        elif re.match(r'^\d+\.\d+\.\d+', line):
            # X.X.X - 三级标题
            add_para(doc, line, size=SIZE_BODY, bold=True, after=6)

        elif is_day_title(line):
            # "第N天（XX条记录）" - 同参考文档的日期行格式
            add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=6)

        elif is_student_profile_header(line):
            # "B-01 | 重点追踪" - 学生个体画像头
            add_para(doc, line, size=SIZE_STUDENT_NAME, bold=True, after=10)

        elif line.startswith(('状态标签', '技能进展', '日志正文', '闪亮一刻',
                             '观察笔记', '综合标注', '筛选结论', '数据汇总')):
            # 标签行 - 与参考文档一致，15pt Bold
            if '：' in line or '：' in line:
                sep = '：' if '：' in line else '：'
                label, value = line.split(sep, 1)
                add_label_para(doc, label + sep, value, size=SIZE_SUB_HEADING, label_bold=True)
            else:
                add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=4)

        elif line.startswith(('访谈ID', '访谈类型', '时长', '内容摘要', '主题标签',
                             '归档人', '归档时间', '访谈转录', '访谈记录', '访谈资料',
                             '创作动机', '绘画经验', '是否重点', '每日追踪',
                             '年级：', '性别：', '编号：', '日期：',
                             '关键词', '学生：', '摘要')):
            # 字段标签行
            if '：' in line or '：' in line:
                sep = '：' if '：' in line else '：'
                label, value = line.split(sep, 1)
                add_label_para(doc, label + sep, value, size=SIZE_BODY, label_bold=True)
            else:
                add_para(doc, line, size=SIZE_BODY, bold=True, after=3)

        elif is_stats_line(line):
            # 统计摘要行
            add_para(doc, line, size=SIZE_BODY, after=3)

        elif re.match(r'^[（(]第\d+天[）)]\s*$', line):
            # 如 "(第1天)"
            add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=6)

        elif re.match(r'^[A-Z]-\d+\s+[｜|]', line):
            # 学生信息行
            add_para(doc, line, size=SIZE_BODY, after=3)

        elif re.match(r'^#\s+\d+', line):
            # GSES题号行
            add_para(doc, line, size=SIZE_BODY, bold=True, after=2)

        elif re.match(r'^\d+\s+\S', line) and len(line.split()) >= 3:
            # GSES数据行
            add_para(doc, line, size=SIZE_BODY, after=2)

        elif re.match(r'^关键事件\d*[：:]\s*$', line):
            add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=6)

        elif re.match(r'^关注维度', line):
            add_para(doc, line, size=SIZE_SUB_HEADING, bold=True, after=6)

        elif line.startswith('学生人数') or line.startswith('观察记录') or \
             line.startswith('重点追踪') or line.startswith('关键事件') or \
             line.startswith('访谈资料') or line.startswith('GSES') or \
             line.startswith('成长日志') or line.startswith('课程名称') or \
             line.startswith('观察者') or line.startswith('课程开始') or \
             line.startswith('导出时间'):
            # 汇总统计行
            if '：' in line or '：' in line:
                sep = '：' if '：' in line else '：'
                label, value = line.split(sep, 1)
                add_label_para(doc, label + sep, value, size=SIZE_BODY, label_bold=True)
            else:
                add_para(doc, line, size=SIZE_BODY, after=4)

        elif re.match(r'^[A-Z]-\d+\s', line) and len(line.split()) >= 6:
            # 观察记录的单行数据（如 "B-01 上午勾线 刻版 持续专注..."）
            add_para(doc, line, size=SIZE_BODY, after=2)

        elif re.match(r'^[（(]\d+[）)]\s*$', line):
            pass  # 纯页码

        elif re.match(r'^\d+\s*/\s*\d+\s*$', line):
            pass  # 页码

        else:
            # 默认正文
            add_para(doc, line, size=SIZE_BODY, after=3)

        li += 1

    # 页间分页
    if page_idx < total_pages - 1:
        doc.add_page_break()

# ============================================================
# 保存
# ============================================================

print("Saving optimized Word document...")
doc.save(OUTPUT_PATH)
size_kb = os.path.getsize(OUTPUT_PATH) / 1024
print(f"[OK] Done! File: {OUTPUT_PATH} ({size_kb:.1f} KB)")
